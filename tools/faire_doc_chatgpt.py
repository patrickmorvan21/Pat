#!/usr/bin/env python3
"""
Génère le .docx de compte rendu à transmettre à ChatGPT.

Le texte vit ICI, en clair, pour rester diffable — le .docx n'est qu'un
emballage. Regénérer : python3 tools/faire_doc_chatgpt.py [sortie.docx]

Convention d'écriture (le lecteur est une IA qui doit pouvoir VÉRIFIER) :
identifiants réels, chiffres mesurés, et les manques dits en clair. Aucune
formule de politesse, aucun superlatif.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ORANGE = RGBColor(0xE0, 0x63, 0x2A)
CHARBON = RGBColor(0x1C, 0x1A, 0x16)

TEXTE = """
# PACTUM — chantier « FEEDBACK + FLUIDITÉ »
## Rapport d'exécution — 12 août 2026, build v1.79.0

Réponse au mémo « PACTUM — CHANTIER FEEDBACK + FLUIDITÉ » (14 sections).
Tout ce qui suit est en ligne et vérifiable en jouant :
https://patrickmorvan21.github.io/Pat/aldenhar/

Le mémo est traité dans son ORDRE D'EXÉCUTION (§12) : audit global d'abord,
propositions ensuite, code après, lot contrôlé et mesure avant toute
généralisation. Rien n'a été élargi sans mesure.

---

# 1. L'audit d'abord (§12.1 à §12.6)

Six relevés produits avant la moindre ligne de code, par un outil versionné
(`tools/audit_fluidite.py`) plus des vérifications DOM en jeu. Trois de ces
relevés ont CHANGÉ le diagnostic du mémo. Ils sont donnés ici parce qu'ils
expliquent pourquoi certaines corrections ne sont pas celles qui étaient
attendues.

**§12.2 — les scènes à plus d'un tap de lecture.**
Sur la narration ÉCRITE : 0 scène sur 80 dépasse un tap. Le chiffre est vrai
et il est trompeur. À l'exécution, l'arrivée dans un lieu empile des blocs que
la scène n'écrit pas : phrase d'approche, franchissement du village, puce
Jour, carte d'état, rumeur, manifestation du Soupçon, ligne de mémoire d'un
PNJ — jusqu'à douze injections possibles sur un même écran.

**Conséquence : une part du levier n'est pas l'écriture, c'est l'ASSEMBLAGE.**
Aucune des lignes lourdes mesurées (serment-hameau 2,7 taps ; marche-muet
1,8 ; tour-de-guet 1,5) n'est corrigeable en raccourcissant un paragraphe.

**⚠️ CORRECTION — la première version de ce rapport concluait de là qu'il ne
fallait pas raccourcir les textes. C'était faux, et sur deux plans.** D'abord
le chiffre : j'avançais « la médiane est déjà à 55 mots », qui est le total
d'une SCÈNE, tous paragraphes confondus — pas ce que le joueur voit. Mesurée
sur des vies réelles, la médiane par ÉCRAN est de **34 mots**. Ensuite le
raisonnement : « raccourcir ne fait pas gagner un tap » répond à une question
de fluidité, alors que la demande portait sur le RYTHME. Les deux ne sont pas
la même chose. Le §7 ci-dessous donne la mesure et la cible corrigée.

**§12.4 — les CTA tronqués.** Un seul CTA est réellement tronqué en jeu
(vérifié au DOM, pas à la longueur) : « Pourquoi les pointes vers
l'intérieur ? » chez le Veilleur, qui déborde de 38 px.

**Et la cause n'était pas sa longueur.** Il débordait à 14 px, c'est-à-dire
sans que le composant de rétrécissement (`FitLabel`) ait réduit quoi que ce
soit. Le composant ajustait la police dans un `useLayoutEffect`, or la barre
de choix est en `display: none` pendant toute la frappe du texte : un élément
dans un parent masqué a `clientWidth === 0`, la condition `scrollWidth >
clientWidth` était donc fausse, et la boucle de réduction ne s'exécutait
jamais. **Le rétrécisseur était mort sur TOUS les boutons du jeu.** Un seul
libellé était assez long pour que ça se voie, mais le défaut était général et
serait revenu à chaque futur libellé un peu long.

**§12.5 — la logique d'entrée / sortie du Hameau.** Le mémo décrit « le joueur
peut entrer deux fois dans le Hameau ». Les quatre garde-fous attendus
existaient déjà (le pool exclut un lieu visité, la séquence d'arrivée
disparaît après l'entrée, il y a un beat de franchissement, un drapeau marque
la halte). Le défaut était réel mais venait d'ailleurs.

Sur 17 destinations tirables, **cinq sont à l'INTÉRIEUR du village** (Marché
Muet, Tour de Guet, Chapelle des Cordes, Puits Condamné, Petit Tribunal). Une
fois entré, le générateur ne retirait que la PORTE : ces cinq lieux restaient
des destinations ordinaires, mélangées aux lieux de lande. On sortait du
village, on marchait dans la bruyère, et la Croisée suivante proposait la
Chapelle — qui est dans le village.

**Ce n'est pas « deux entrées », c'est une absence de frontière.**

**§12.6 — le tableau Action → Feedback → Effet → Consommateur.** Aucune action
importante n'est muette : les 100+ choix qui posent un effet ont tous une
conséquence écrite ou les quatre issues d'un jet. Le §7 était donc déjà tenu
sur le texte. En revanche, **huit actions posaient un drapeau que personne ne
lisait** — dont un posé la veille par moi-même, avec la promesse écrite dans
le code qu'il ouvrirait une option. Il ne l'ouvrait pas. C'est exactement ce
que le §12 appelle « un flag stocké pris pour une fonctionnalité terminée ».

---

# 2. Ce qui a été livré

Cinq correctifs, deux livraisons (v1.78.0 puis v1.79.0), chacune déployée et
prouvée en jeu.

## Correctif 1 — le rétrécisseur de libellés réparé (§6, §13, §14)

L'ajustement se refait désormais chaque fois que l'élément acquiert une
largeur réelle (`ResizeObserver`), c'est-à-dire au moment précis où la barre
repasse en `flex`. Prouvé en jeu : plus aucun libellé en débordement, et les
polices apparaissent enfin à 12 / 12,5 / 14 px là où elles étaient toutes
bloquées au maximum. Le critère « 0 CTA essentiel tronqué » est refermé pour
de bon, plutôt que raccourci au cas par cas avec le piège laissé armé.

## Correctif 2 — le village devient une enclave (§3, §13)

Trois états au lieu de deux : dehors, dedans, sorti. Tant qu'on est dedans, la
Croisée ne propose que des lieux du village et garantit toujours une issue ;
une fois sorti (drapeau `hameau.sorti`), ni le village ni ses rues ne sont
retirables de la vie. Les deux garanties du système (le chapitre du Bailli, la
Colline) ont été réalignées dessus — sans quoi elles réintroduisaient le
village après coup.

Prouvé en jeu dans les deux sens : depuis une rue, la Croisée offre encore
`petit-tribunal` ; une fois sorti, elle n'offre que des lieux de lande, sans
rue et sans porte.

## Correctif 3 — un seul rappel par arrivée (§5)

C'est la réponse au vrai levier identifié au §12.2. Les injections d'arrivée
passent par un budget : familiarité, mémoire d'un PNJ, réaction d'état,
perception, mode d'arrivée — un seul passe, par priorité, et lui seul consomme
son effet de bord (les autres ne sont pas « brûlés » et resserviront).

Mesuré sur une arrivée réelle au Champ des Fixés, compte vétéran :
**6 → 5 taps et 3 → 2 blocs de narration.**

## Correctif 4 — l'objet agit sur place (§2, §7, §9)

« L'objet doit d'abord modifier la scène actuelle, PUIS ouvrir ou transformer
une possibilité. » Auparavant, sortir un objet appliquait son effet et
enchaînait l'écran suivant : le joueur ne voyait jamais ce que son objet
venait de changer — la « récompense invisible » du §9.

La branche d'usage passe désormais TOUJOURS par « rester sur place ». La
conséquence s'écrit, l'écran se recompose, et ce que l'objet ouvre apparaît
dedans. Trois champs de données, de portée ÉCRAN : la scène déclare l'objet
qui sert chez elle (`usageObjet`), un choix peut exiger l'usage
(`requiresUsage`) ou disparaître avec lui (`masqueSiUsage`). Le jeton transite
par le registre des choix faits, vidé en quittant le lieu : aucune donnée
persistée de plus, et une corde amarrée ici n'ouvre rien ailleurs. C'est la
règle, pas une limite — l'objet transforme la scène, pas la partie.

**Objet pilote n°1, l'OUTIL — la Corde coupée.** Le premier objet des Landes
qui ne soigne rien et ne pèse sur aucun jet : il ouvre un endroit. Elle se
gagne enfin à la Chapelle des Cordes (le choix « Prendre la corde coupée »
disait depuis toujours qu'elle s'enroule autour du poignet, et on repartait
les mains vides) et s'amarre à la margelle du Puits Condamné. Les coups
s'arrêtent, la corde file entre deux planches, puis quelque chose l'amarre par
l'autre bout. Alors « Descendre par la corde » apparaît.

Le budget de trois CTA tient **sans exception** : une fois la corde tendue par
le fond, coller l'oreille aux planches n'est plus la question — les deux
options d'observation disparaissent. Mesuré : **2 CTA après usage.**

**Objet pilote n°2, le SOIN — le baume.** Il ne saute plus l'écran. La ligne
dit ce qui se referme ET ce que ça change pour la suite (« Ce que tu vas
tenter maintenant, tu le tenteras avec ce corps-là »), sans un chiffre :
l'Anneau du prochain jet montre la différence de lui-même, puisque le malus de
blessure vient de tomber. C'est le volet « transformer une possibilité ».

**Deux objets seulement, délibérément** — le modèle est éprouvé avant d'être
étendu aux dix autres, conformément au §12.8.

## Correctif 5 — les huit effets orphelins (§5, §8, §9)

Trois lecteurs, cinq coupes.

Les trois lecteurs portent tous l'arc du Bailli et de sa fille, et illustrent
le §8 (« le jeu a retenu ça ») :

— La confession du Pendu qui parle (« j'ai signé trois cents noms, le trois
cent unième était le sien, alors j'ai inscrit le mien en dessous ») devient
une DÉFENSE à ton propre procès. Le hameau juge par l'Ordonnance de la
Fixation, et l'homme qui l'a écrite s'est condamné pour avoir refusé de
l'appliquer. Le joueur ne l'apprend pas au procès : il l'a apprise ailleurs,
peut-être dans une vie précédente, et elle revient quand il en a besoin.

— Ce que la Fille du Moulin nomme sur le toit de la grange (« celle qui
regarde — elle attend que vous le fassiez ») ouvre « Ne rien lui donner » : la
seule chose que la chose n'attendait pas.

— Les combles clouées de l'intérieur se comptent à l'aube, en sortant. Le
nombre dit combien de familles se sont barricadées la même nuit.

Les cinq autres drapeaux sont COUPÉS : la prose reste intacte, seul le drapeau
mort part. Les garder « pour un Codex plus tard » est précisément ce qui a
rendu cet audit nécessaire ; re-poser une ligne quand le Codex existera coûte
une ligne.

Après coup, l'audit rend : **0 action sans feedback, 0 effet sans
consommateur.**

---

# 3. Les garde-fous ajoutés

Le mémo insiste (§9, §12, §13) sur le fait qu'un drapeau qui fonctionne n'est
pas une fonctionnalité. Deux mécanismes empêchent la rechute, et ils cassent
le build plutôt que d'attendre un playtest.

**Garde `A-usage`.** Une scène qui déclare un usage d'objet doit avoir un
choix du MÊME écran qui lit sa clé ; l'objet doit exister au catalogue ; un
choix conditionné à un usage qu'aucune scène ne déclare est signalé. Sans
quoi on paierait un objet rare pour une ligne de texte, ou une option
n'apparaîtrait jamais quoi que fasse le joueur.

**Le garde a été prouvé sur ses trois modes d'échec AVANT d'être branché** —
règle de méthode du projet : un garde qu'on n'a pas vu tirer n'est pas un
garde. (Quatre extracteurs d'audit se sont déjà révélés muets sur ce projet :
ils rendaient zéro signalement parce qu'ils ne lisaient rien.)

Le build tourne désormais derrière sept gardes, tous verts sur cette
livraison.

---

# 4. Ce qui a été mesuré, et ce que la mesure ne dit pas

Banc instrumenté (`tools/mesure_boucle.mjs`), deux profils de joueur, cinq
métriques, définition stricte de « décision » : ouvrir un sous-menu ou
examiner un point d'intérêt N'EST PAS une décision. Sans cette règle, la
mesure se retournerait contre le chantier — supprimer les sous-menus ferait
BAISSER le score.

Les trois critères bloquants passent :

— aucun sous-menu générique ouvert sur le lot (0 sur 11 auparavant) ;

— au plus 3 décisions visibles à l'arrivée (maximum observé : 3) ;

— moins d'écrans avant décision : 1,98 → 1,76 tap.

**Ce que ces chiffres NE prouvent PAS, et il faut le dire.** Le banc mesure
trois lieux, dont aucun n'est touché par les correctifs 4 et 5. Les écarts du
tableau sont du bruit de tirage, pas un effet. La preuve du bruit est
disponible : deux passages du banc sur le MÊME build, même lieu, même profil,
ont donné 33 % / 2,0 tap et 38 % / 1,6 tap. Le banc sert à détecter une
régression franche, pas à départager deux points de pourcentage.

**Deuxième limite, structurelle :** le banc pose le héros directement dans un
lieu, donc **il ne joue jamais d'arrivée**. Il ne peut pas mesurer le gain du
correctif 3, qui porte précisément sur l'assemblage d'arrivée — c'est-à-dire
sur la cause identifiée des 4-5 taps ressentis. Une mesure qui porterait sur
ce ressenti demande des vies entières, pas des lieux isolés. C'est un chantier
d'outillage d'environ une session, non engagé faute d'arbitrage.

**Ce qui est prouvé, l'est en jeu et pas au banc :** 29 assertions sur 29,
dans des runs réelles, avec drags du dé — l'outil offert, la descente absente
avant l'usage puis présente, l'écran inchangé, l'objet consommé, les trois
lecteurs absents sans leur découverte et présents avec, et le compte de CTA à
chaque étape.

---

# 5. Ce qui n'a PAS été fait, et pourquoi

**Les 29 points d'intérêt restants, sur 11 scènes.** Le §12.8 demande
d'implémenter sur un lot contrôlé, de mesurer, puis de généraliser. Le lot
pilote (Colline, Moulin, Hameau) a été traité et mesuré ; l'extension aux 11
autres scènes attend un arbitrage. Le libellé générique « Observer les
alentours » n'existe plus que comme conséquence : il n'apparaît que si une
scène a des points. **Supprimer les points supprime le bouton** — il n'y a
donc aucun sous-menu résiduel à traquer, seulement du contenu à convertir.

**La cible de 25-45 mots (§10).** Elle n'est pas appliquée à ce jour — mais
elle n'est plus refusée : la mesure sur vies complètes (§7) montre qu'elle est
le chantier suivant, avec une cible plus précise que « raccourcir tout ».

**Le Codex.** Écran non conçu, maquettes attendues. Les cinq drapeaux coupés
en §2 le concernaient ; ils seront reposés à ce moment-là, une ligne chacun.

---

# 6. Points ouverts

**Une icône partagée.** La Corde coupée et la Mèche Nouée utilisent le même
fichier. L'image montre un rouleau de chanvre épais : elle a été nommée pour
la corde et servait à une mèche de cheveux. Elle a été rendue à la corde ; la
mèche l'emprunte en attendant la sienne.

**Un piège d'infrastructure, corrigé, qui vaut d'être connu.** Le worktree de
déploiement est en HEAD détaché : un `git push -u origin gh-pages` y pousse la
branche locale périmée et répond « Everything up-to-date ». Le rituel paraît
s'être bien déroulé et rien n'est en ligne. Il est possible que des
déploiements passés en aient souffert sans qu'on le voie. La parade est
inscrite au journal du projet : pousser explicitement la référence et vérifier
l'égalité des empreintes avant de retirer le worktree.

**Question de méthode, pour la suite.** Le mémo demandait de mesurer « le
nombre de taps avant la prochaine décision significative » (§5). C'est la
bonne métrique, et c'est celle que l'outillage actuel mesure le moins bien,
puisqu'il ne joue pas d'arrivée. Faut-il investir la session d'outillage qui
permettrait de la mesurer sur des vies entières, ou s'en tenir au jugement en
playtest ?

---

# 7. LE TEST DES DEUX VIES — la mesure qui manquait

Deux vies COMPLÈTES rejouées sur v1.79, une curieuse et une pressée, sur le
build publié : 167 écrans, 53 décisions, 0 erreur. C'est précisément ce que le
banc ne savait pas faire — il posait le héros dans un lieu et ne jouait jamais
d'arrivée. Nouvel outil : `tools/rythme.py`, qui lit les transcripts de
parties réelles. Une base de comparaison existe : deux vies enregistrées de la
même façon sur v1.59, avant tout ce chantier.

**Taps de lecture avant une décision (la métrique du §5) :**

v1.59 → **1,98** tap · 36 % des décisions à un tap ou moins

v1.79 → **1,79** tap · 38 % des décisions à un tap ou moins

**Verdict : la cible du §10 n'est PAS atteinte.** « La majorité des scènes
mène à une décision après au plus un tap » demande plus de 50 % ; on est à
38 %. Le gain sur la moyenne est réel mais mince, et la part des décisions
immédiates n'a quasiment pas bougé. Sur la vie pressée, c'est pire : **19 %**
seulement, et une décision sur trois demande trois taps.

**Longueur des écrans — et c'est ici que la critique du §10 était juste :**

médiane **34 mots** (bonne, et à ne pas toucher)

mais **46 % des écrans dépassent 45 mots**, contre 42 % sur v1.59

et la queue est lourde : 149, 108, 106, 96, 94 mots sur un seul écran

**Ce n'est donc pas « tout raccourcir » qu'il faut, c'est COUPER LA QUEUE.**
La médiane est déjà dans la cible ; ce sont les 46 % au-dessus de 45 mots, et
surtout la poignée d'écrans à 90-150 mots, qui coûtent le tap
supplémentaire. C'est une cible mesurable, et `tools/rythme.py` la vérifie.

**Où l'on tapote sans décider.** L'outil nomme les endroits où trois taps ou
plus s'enchaînent sans décision. Les voici, et ils sont instructifs : ce sont
tous des écrans de TEXTE, jamais des menus.

« Ils ont commencé à te compter. C'est un pays méticuleux. » · « Tu ne sens
rien ? Normal. On ne sent jamais le premier tour de corde. » · « Un seul
écriteau détonne : le motif a été gratté au couteau… » · « Ça grince en
mesure. Toutes ces cordes, un seul rythme. »

Autrement dit : les sous-menus et l'assemblage d'arrivée ont bien été traités
— ce qui reste, c'est la prose. La critique était fondée.

**Les autres questions posées, telles que les vies y répondent :**

**« Observer » manque-t-il ?** Non — mais il est encore là : 7 apparitions
dans la vie curieuse, sur les scènes hors lot pilote. Les 29 points restants
se voient donc réellement en jeu.

**Les objets récompensent-ils ?** Partiellement. 3 objets gagnés par vie, mais
1 seul sorti sur la vie curieuse et 0 sur la pressée. Le nouveau geste
(l'objet qui transforme la scène) fonctionne quand il est déclenché, mais il
est encore trop rare pour porter une sensation de récompense. Deux objets
pilotes sur douze : c'est attendu, et ça confirme qu'il faut étendre.

**Les choix reviennent-ils hanter ?** Oui, et c'est le point le plus solide :
9 à 10 échos de mémoire par vie, dont des reconnaissances explicites de
personnages (« Je t'ai déjà croisé »). La chaîne « j'ai fait quelque chose →
le monde l'a retenu » est en place et se voit sans qu'on la cherche.

**Ce qui ne s'est pas produit du tout :** aucun Destin ni Malédiction sur les
deux vies, et le Geôlier n'a parlé qu'une fois en 167 écrans. Les moments les
plus forts du jeu sont donc rares au point d'être absents d'une partie
entière — c'est un sujet de dosage distinct, à regarder ensuite.

**Conclusion de ce test : le feu vert ne doit pas être donné.** Non pas parce
que les correctifs livrés seraient faux — ils tiennent, et la mémoire du monde
fonctionne — mais parce que la métrique qui décide, elle, n'est pas atteinte :
38 % au lieu d'une majorité. Le chantier suivant est la queue des textes, pas
la généralisation des 29 points. Les deux sont à faire ; celui-là d'abord.

---

# 8. Définition de « fini » (§13) — état point par point

**Un objet utilisé montre ce qu'il modifie** — fait, sur deux objets pilotes,
prouvé en jeu.

**Une action récompensée a au moins un signal immédiat** — fait, 0 action
muette à l'audit.

**Le Hameau ne peut pas être ré-entré après sa sortie** — fait, prouvé dans
les deux sens.

**Aucun bouton « Observer » générique dans la boucle** — partiel : supprimé du
lot pilote, 29 points restants sur 11 scènes, en attente d'arbitrage.

**Aucun CTA essentiel tronqué** — fait, et la cause systémique est réparée,
pas contournée.

**La majorité des scènes mène à une décision après au plus un tap** — NON
ATTEINT, désormais mesuré : 38 % sur deux vies complètes (§7). C'est le
critère qui bloque le feu vert.
""".strip()


def construire(sortie: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = CHARBON

    # ⚠️ Découpage par LIGNE, pas par bloc : deux titres consécutifs (le titre
    # et son sous-titre) vivent dans le même bloc et fusionnaient en un seul
    # paragraphe. On ferme donc le paragraphe courant dès qu'une ligne est un
    # titre ou un séparateur.
    blocs: list[str] = []
    courant: list[str] = []

    def fermer() -> None:
        if courant:
            blocs.append(" ".join(courant))
            courant.clear()

    for ligne in TEXTE.split("\n"):
        ligne = ligne.rstrip()
        if not ligne:
            fermer()
        elif ligne.startswith(("# ", "## ")) or ligne == "---":
            fermer()
            blocs.append(ligne)
        else:
            courant.append(ligne.strip())
    fermer()

    for bloc in blocs:
        if bloc == "---":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run("• • •")
            r.font.color.rgb = ORANGE
            continue
        if bloc.startswith("# "):
            h = doc.add_heading(bloc[2:].strip(), level=1)
            for r in h.runs:
                r.font.color.rgb = ORANGE
            continue
        if bloc.startswith("## "):
            h = doc.add_heading(bloc[3:].strip(), level=2)
            for r in h.runs:
                r.font.color.rgb = CHARBON
            continue
        # Un paragraphe ordinaire : le gras **…** est le seul balisage inline.
        p = doc.add_paragraph()
        for i, part in enumerate(bloc.split("**")):
            if not part:
                continue
            run = p.add_run(part)
            run.bold = i % 2 == 1
    doc.save(sortie)


if __name__ == "__main__":
    dest = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(
        "data/PACTUM-feedback-fluidite-rapport.docx")
    dest.parent.mkdir(parents=True, exist_ok=True)
    construire(dest)
    print(f"écrit : {dest} ({dest.stat().st_size} octets)")
